import tkinter as tk
from tkinter import messagebox
from PyQt5.QtWidgets import QApplication, QWidget, QLabel, QPushButton, QGridLayout, QFileDialog, QMessageBox, QInputDialog
from PyQt5.QtCore import QDir, Qt
from PIL import Image, ImageQt
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from pydub import AudioSegment
from pdf2docx import Converter
from PyPDF2 import PdfMerger, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor

class FileConverterApp(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()
    
    def initUI(self):
        self.setWindowTitle('IMS Ghaziabad Private File Conversion Tool')
        self.setGeometry(100, 100, 400, 300)
        self.setStyleSheet("background-color: #202243; color: white;")

        # Add header label
        self.lbl_header = QLabel('IMS Ghaziabad Private File Conversion Tool', self)
        self.lbl_header.setStyleSheet("color: white; font-size: 16pt; font-weight: bold;")
        self.lbl_header.setAlignment(Qt.AlignCenter)

        self.lbl_info = QLabel('', self)
        self.lbl_info.setStyleSheet("color: white;")

        self.lbl_image = QLabel('Image', self)
        self.lbl_image.setStyleSheet("color: white;")
        self.lbl_audio = QLabel('Audio', self)
        self.lbl_audio.setStyleSheet("color: white;")
        self.lbl_document = QLabel('Document', self)
        self.lbl_document.setStyleSheet("color: white;")
        self.lbl_utility = QLabel('Utility', self)
        self.lbl_utility.setStyleSheet("color: white;")

        self.btn_png_to_jpg = QPushButton('PNG to JPG', self)
        self.btn_jpg_to_png = QPushButton('JPG to PNG', self)
        self.btn_png_to_pdf = QPushButton('PNG to PDF', self)
        self.btn_mp3_to_wav = QPushButton('MP3 to WAV', self)
        self.btn_pdf_to_doc = QPushButton('PDF to DOCX', self)
        self.btn_doc_to_pdf = QPushButton('DOCX to PDF', self)
        self.btn_encrypt = QPushButton('Encrypt', self)
        self.btn_decrypt = QPushButton('Decrypt', self)
        self.btn_webp_to_png = QPushButton('WebP to PNG', self)
        self.btn_png_to_ico = QPushButton('PNG to ICO', self)
        self.btn_doc_to_txt = QPushButton('DOCX to TXT', self)
        self.btn_merge_pdf = QPushButton('Merge PDF', self)
        self.btn_lock_pdf = QPushButton('Lock PDF', self)
        self.btn_pdf_to_csv = QPushButton('PDF to CSV', self)  # Added button

        self.btns = [self.btn_png_to_jpg, self.btn_jpg_to_png, self.btn_png_to_pdf, self.btn_mp3_to_wav,
                     self.btn_pdf_to_doc, self.btn_doc_to_pdf, self.btn_encrypt, self.btn_decrypt,
                     self.btn_webp_to_png, self.btn_png_to_ico, self.btn_doc_to_txt, self.btn_merge_pdf,
                     self.btn_lock_pdf, self.btn_pdf_to_csv]  # Added button

        for btn in self.btns:
            btn.setStyleSheet("background-color: #2d325a; color: white; border: none; border-radius: 10px;")
        
        self.layout = QGridLayout()
        self.layout.addWidget(self.lbl_header, 0, 0, 1, 2)
        self.layout.addWidget(self.lbl_info, 1, 0, 1, 2)

        # Image section
        self.layout.addWidget(self.lbl_image, 2, 0)
        self.layout.addWidget(self.btn_png_to_jpg, 3, 0)
        self.layout.addWidget(self.btn_jpg_to_png, 3, 1)
        self.layout.addWidget(self.btn_png_to_pdf, 4, 0)
        self.layout.addWidget(self.btn_webp_to_png, 4, 1)
        self.layout.addWidget(self.btn_png_to_ico, 5, 0)

        # Audio section
        self.layout.addWidget(self.lbl_audio, 7, 0)
        self.layout.addWidget(self.btn_mp3_to_wav, 8, 0)

        # Document section
        self.layout.addWidget(self.lbl_document, 10, 0)
        self.layout.addWidget(self.btn_pdf_to_doc, 11, 0)
        self.layout.addWidget(self.btn_doc_to_pdf, 11, 1)
        self.layout.addWidget(self.btn_doc_to_txt, 12, 0)
        self.layout.addWidget(self.btn_merge_pdf, 12, 1)

        # Utility section
        self.layout.addWidget(self.lbl_utility, 14, 0)
        self.layout.addWidget(self.btn_encrypt, 15, 0)
        self.layout.addWidget(self.btn_decrypt, 15, 1)
        self.layout.addWidget(self.btn_lock_pdf, 16, 0)
        self.layout.addWidget(self.btn_pdf_to_csv, 16, 1)  # Added button

        # Completion message label
        self.lbl_completion = QLabel('', self)
        self.lbl_completion.setStyleSheet("color: white;")
        self.layout.addWidget(self.lbl_completion, 19, 0, 1, 2)

        self.setLayout(self.layout)

        # Connect button clicks to conversion methods
        self.btn_png_to_jpg.clicked.connect(lambda: self.conversion_window('PNG to JPG', 'jpg'))
        self.btn_jpg_to_png.clicked.connect(lambda: self.conversion_window('JPG to PNG', 'png'))
        self.btn_png_to_pdf.clicked.connect(lambda: self.conversion_window('PNG to PDF', 'pdf'))
        self.btn_mp3_to_wav.clicked.connect(lambda: self.conversion_window('MP3 to WAV', 'wav'))
        self.btn_pdf_to_doc.clicked.connect(lambda: self.conversion_window('PDF to DOCX', 'docx'))
        self.btn_doc_to_pdf.clicked.connect(lambda: self.conversion_window('DOCX to PDF', 'pdf'))
        self.btn_webp_to_png.clicked.connect(lambda: self.conversion_window('WebP to PNG', 'png'))
        self.btn_png_to_ico.clicked.connect(lambda: self.conversion_window('PNG to ICO', 'ico'))
        self.btn_doc_to_txt.clicked.connect(lambda: self.conversion_window('DOCX to TXT', 'txt'))
        self.btn_merge_pdf.clicked.connect(self.merge_pdf_files)
        self.btn_lock_pdf.clicked.connect(self.lock_pdf_window)
        self.btn_encrypt.clicked.connect(self.encrypt_file_window)
        self.btn_decrypt.clicked.connect(self.decrypt_file_window)
        self.btn_pdf_to_csv.clicked.connect(lambda: self.conversion_window('PDF to CSV', 'csv'))  # Added button

    def conversion_window(self, conversion_type, output_extension):
        input_file, _ = QFileDialog.getOpenFileName(self, f'Select Input File for {conversion_type}', '', 'All Files (*)')
        if input_file:
            output_folder = QFileDialog.getExistingDirectory(self, 'Select Output Folder', QDir.homePath())
            if output_folder:
                try:
                    if conversion_type == 'MP3 to WAV':
                        audio = AudioSegment.from_file(input_file)
                        output_file = os.path.join(output_folder, f'{os.path.splitext(os.path.basename(input_file))[0]}.{output_extension}')
                        audio.export(output_file, format=output_extension)
                    elif conversion_type == 'PDF to DOCX':
                        output_file = os.path.join(output_folder, f'{os.path.splitext(os.path.basename(input_file))[0]}.{output_extension}')
                        self.convert_pdf_to_doc(input_file, output_file)
                    elif conversion_type == 'DOCX to PDF':
                        output_file = os.path.join(output_folder, f'{os.path.splitext(os.path.basename(input_file))[0]}.{output_extension}')
                        self.convert_doc_to_pdf(input_file, output_file)
                    elif conversion_type == 'WebP to PNG':
                        output_file = os.path.join(output_folder, f'{os.path.splitext(os.path.basename(input_file))[0]}.{output_extension}')
                        self.convert_webp_to_png(input_file, output_file)
                    elif conversion_type == 'PNG to ICO':
                        output_file = os.path.join(output_folder, f'{os.path.splitext(os.path.basename(input_file))[0]}.{output_extension}')
                        self.convert_png_to_ico(input_file, output_file)
                    elif conversion_type == 'DOCX to TXT':
                        output_file = os.path.join(output_folder, f'{os.path.splitext(os.path.basename(input_file))[0]}.{output_extension}')
                        self.convert_docx_to_txt(input_file, output_file)
                    elif conversion_type == 'PDF to CSV':  # Added condition
                        output_file = os.path.join(output_folder, f'{os.path.splitext(os.path.basename(input_file))[0]}.{output_extension}')
                        self.convert_pdf_to_csv(input_file, output_file)  # Added method
                    else:
                        img = Image.open(input_file)
                        if conversion_type == 'JPG to PNG':
                            img = img.convert('RGBA')
                        elif img.mode == 'RGBA':
                            img = img.convert('RGB')
                        file_name_without_extension = os.path.splitext(os.path.basename(input_file))[0]
                        output_file = os.path.join(output_folder, f'{file_name_without_extension}.{output_extension}')
                        img.save(output_file)
                    self.lbl_completion.setText(f'{conversion_type} conversion successful. Output file: {output_file}')
                except Exception as e:
                    QMessageBox.critical(self, 'Error', f'Error converting file: {e}')
            else:
                QMessageBox.warning(self, 'Warning', 'Output folder not selected.')
        else:
            QMessageBox.warning(self, 'Warning', 'Input file not selected.')
    
    def convert_pdf_to_doc(self, input_file, output_file):
        try:
            cv = Converter(input_file)
            cv.convert(output_file, start=0, end=None)
            cv.close()
        except Exception as e:
            raise e
    
    def convert_doc_to_pdf(self, input_file, output_file):
        try:
            doc = Document(input_file)
            styles = getSampleStyleSheet()
            pdf = SimpleDocTemplate(output_file, pagesize=letter)

            story = []

            for para in doc.paragraphs:
                pdf_para = Paragraph(para.text, styles['Normal'])
                story.append(pdf_para)

            pdf.build(story)
        except Exception as e:
            raise e

    def convert_webp_to_png(self, input_file, output_file):
        try:
            img = Image.open(input_file)
            img.save(output_file, 'png')
        except Exception as e:
            raise e

    def convert_png_to_ico(self, input_file, output_file):
        try:
            img = Image.open(input_file)
            img.save(output_file, 'ico')
        except Exception as e:
            raise e

    def convert_docx_to_txt(self, input_file, output_file):
        try:
            doc = Document(input_file)
            with open(output_file, 'w') as txt_file:
                for para in doc.paragraphs:
                    txt_file.write(para.text)
                    txt_file.write('\n')
        except Exception as e:
            raise e

    def convert_pdf_to_csv(self, input_file, output_file):  # Added method
        try:
            # Add logic to convert PDF to CSV here
            pass
        except Exception as e:
            raise e

    def merge_pdf_files(self):
        input_files, _ = QFileDialog.getOpenFileNames(self, 'Select PDF Files to Merge', '', 'PDF Files (*.pdf)')
        if input_files:
            output_file, _ = QFileDialog.getSaveFileName(self, 'Save Merged PDF As', '', 'PDF Files (*.pdf)')
            if output_file:
                try:
                    merger = PdfMerger()
                    for pdf_file in input_files:
                        merger.append(pdf_file)
                    merger.write(output_file)
                    merger.close()
                    self.lbl_completion.setText(f'PDF files merged successfully. Output file: {output_file}')
                except Exception as e:
                    QMessageBox.critical(self, 'Error', f'Error merging PDF files: {e}')
            else:
                QMessageBox.warning(self, 'Warning', 'Output file not selected.')
        else:
            QMessageBox.warning(self, 'Warning', 'No PDF files selected for merging.')

    def lock_pdf_window(self):
        input_file, _ = QFileDialog.getOpenFileName(self, 'Select PDF File to Lock', '', 'PDF Files (*.pdf)')
        if input_file:
            output_folder = QFileDialog.getExistingDirectory(self, 'Select Output Folder', QDir.homePath())
            if output_folder:
                password, ok = QInputDialog.getText(self, 'Enter Password', 'Enter password for PDF:')
                if ok:
                    try:
                        output_file = os.path.join(output_folder, f'{os.path.splitext(os.path.basename(input_file))[0]}_locked.pdf')
                        output_file = self.lock_pdf(input_file, output_file, password)
                        self.lbl_completion.setText(f'PDF locked successfully with password. Output file: {output_file}')
                    except Exception as e:
                        QMessageBox.critical(self, 'Error', f'Error locking PDF: {e}')
            else:
                QMessageBox.warning(self, 'Warning', 'Output folder not selected.')
        else:
            QMessageBox.warning(self, 'Warning', 'No PDF file selected for locking.')

    def lock_pdf(self, input_file, output_file, password):
        try:
            writer = PdfWriter()
            reader = PdfReader(input_file)
            for page in range(len(reader.pages)):
                writer.add_page(reader.pages[page])
            writer.encrypt(user_password=password, owner_password=None, use_128bit=True)
            with open(output_file, 'wb') as f:
                writer.write(f)
            return output_file
        except Exception as e:
            raise e

    def encrypt_file_window(self):
        input_file, _ = QFileDialog.getOpenFileName(self, 'Select File to Encrypt', '', 'All Files (*)')
        if input_file:
            output_folder = QFileDialog.getExistingDirectory(self, 'Select Output Folder', QDir.homePath())
            if output_folder:
                try:
                    encrypted_file_path = self.encrypt_file(input_file)
                    self.lbl_completion.setText(f'File encrypted successfully. Output file: {encrypted_file_path}')
                except Exception as e:
                    QMessageBox.critical(self, 'Error', f'Error encrypting file: {e}')
            else:
                QMessageBox.warning(self, 'Warning', 'Output folder not selected.')
        else:
            QMessageBox.warning(self, 'Warning', 'No file selected for encryption.')

    def decrypt_file_window(self):
        input_file, _ = QFileDialog.getOpenFileName(self, 'Select File to Decrypt', '', 'All Files (*)')
        if input_file:
            output_folder = QFileDialog.getExistingDirectory(self, 'Select Output Folder', QDir.homePath())
            if output_folder:
                try:
                    decrypted_file_path = self.decrypt_file(input_file)
                    output_file = os.path.join(output_folder, os.path.basename(decrypted_file_path))
                    os.rename(decrypted_file_path, output_file)
                    self.lbl_completion.setText(f'File decrypted successfully. Output file: {output_file}')
                except Exception as e:
                    QMessageBox.critical(self, 'Error', f'Error decrypting file: {e}')
            else:
                QMessageBox.warning(self, 'Warning', 'Output folder not selected.')
        else:
            QMessageBox.warning(self, 'Warning', 'No file selected for decryption.')

    def encrypt_file(self, file_path):
        with open(file_path, 'rb') as file:
            content = file.read()
        encrypted_content = bytes([(byte + 1) % 256 for byte in content])  # Shift each byte by 1
        encrypted_file_path = file_path + '.encrypted'
        with open(encrypted_file_path, 'wb') as encrypted_file:
            encrypted_file.write(encrypted_content)
        return encrypted_file_path

    def decrypt_file(self, file_path):
        if file_path.endswith('.encrypted'):
            decrypted_file_path = file_path[:-10]  # Remove the '.encrypted' extension
            with open(file_path, 'rb') as encrypted_file:
                content = encrypted_file.read()
            decrypted_content = bytes([(byte - 1) % 256 for byte in content])  # Reverse the encryption process
            with open(decrypted_file_path, 'wb') as decrypted_file:
                decrypted_file.write(decrypted_content)
            return decrypted_file_path
        else:
            raise ValueError("File is not encrypted")

if __name__ == '__main__':
    app = QApplication([])
    converter_app = FileConverterApp()
    converter_app.show()
    app.exec_()
